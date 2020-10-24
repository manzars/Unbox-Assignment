from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import datetime


def create_offer_letter(user, address, position, start_date, ctc):

    company_name = "Unbox Innovation"
    date = str(datetime.datetime.now()).split(" ")[0]
    user = user
    address = address
    position = position
    start_date = start_date
    ctc = ctc

    document = Document()
    sample_doc = Document('media/demo.docx')

    heading = document.add_heading('OFFER LETTER')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    document.add_paragraph("").add_run(company_name).bold = True
    document.add_paragraph("").add_run(date).bold = True
    document.add_paragraph("").add_run(user).bold = True
    document.add_paragraph("").add_run(address).bold = True

    p = document.add_paragraph(sample_doc.paragraphs[7].runs[0].text)
    p.add_run(user).bold = True
    p.add_run(",")

    p = document.add_paragraph(
        sample_doc.paragraphs[8].runs[0].text + sample_doc.paragraphs[8].runs[1].text)
    p.add_run(company_name).bold = True
    p.add_run(sample_doc.paragraphs[8].runs[3].text)

    p = document.add_paragraph(sample_doc.paragraphs[10].runs[0].text)
    p.add_run(position).bold = True
    p.add_run(sample_doc.paragraphs[10].runs[2].text)
    p.add_run(start_date).bold = True
    p.add_run(sample_doc.paragraphs[10].runs[4].text +
              sample_doc.paragraphs[10].runs[5].text)
    p.add_run(str(ctc) + " Lpa").bold = True
    p.add_run(sample_doc.paragraphs[10].runs[7].text +
              sample_doc.paragraphs[10].runs[8].text)

    p = document.add_paragraph(sample_doc.paragraphs[12].text)

    p = document.add_paragraph(sample_doc.paragraphs[14].runs[0].text)
    p.add_run(company_name).bold = True
    p.add_run(".")

    p = document.add_paragraph(sample_doc.paragraphs[16].text)
    document.add_paragraph("")

    p = document.add_paragraph("")
    p.add_run(user).bold = True

    p = document.add_paragraph("")
    p.add_run("(" + position + ")").bold = True

    document.save('media/' + user.split(" ")[0] + '_offer_letter.docx')
